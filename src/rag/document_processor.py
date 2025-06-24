"""
Document Processing for RAG (Retrieval-Augmented Generation)

This module provides document processing capabilities for the RAG system,
including text extraction, chunking, preprocessing, and metadata extraction
from various document formats.
"""

import asyncio
import logging
import re
from typing import List, Dict, Any, Optional, Union, Tuple
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from enum import Enum
import hashlib
import uuid

# Document processing libraries
import PyPDF2
import docx
import markdown
from bs4 import BeautifulSoup
import json
import yaml

from ..config.settings import Settings
from ..utils.embedding import get_embedding_service


logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """Supported document types"""
    TEXT = "text"
    MARKDOWN = "markdown"
    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"
    JSON = "json"
    YAML = "yaml"
    CODE = "code"
    UNKNOWN = "unknown"


@dataclass
class DocumentChunk:
    """Represents a processed document chunk"""
    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    content: str = ""
    chunk_index: int = 0
    start_char: int = 0
    end_char: int = 0
    metadata: Dict[str, Any] = field(default_factory=dict)
    embedding: Optional[List[float]] = None
    
    def __post_init__(self):
        # Calculate content hash for deduplication
        self.metadata["content_hash"] = hashlib.md5(self.content.encode()).hexdigest()


@dataclass
class ProcessedDocument:
    """Represents a fully processed document"""
    id: str = field(default_factory=lambda: str(uuid.uuid4()))
    title: str = ""
    content: str = ""
    document_type: DocumentType = DocumentType.UNKNOWN
    source_path: Optional[str] = None
    source_url: Optional[str] = None
    chunks: List[DocumentChunk] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    processed_at: datetime = field(default_factory=datetime.utcnow)
    embedding: Optional[List[float]] = None
    
    def get_chunk_count(self) -> int:
        """Get number of chunks"""
        return len(self.chunks)
    
    def get_total_length(self) -> int:
        """Get total content length"""
        return len(self.content)


class TextExtractor:
    """Extracts text from various document formats"""
    
    @staticmethod
    def extract_from_text(content: str) -> str:
        """Extract text from plain text"""
        return content.strip()
    
    @staticmethod
    def extract_from_markdown(content: str) -> str:
        """Extract text from Markdown"""
        try:
            # Convert markdown to HTML then extract text
            html = markdown.markdown(content)
            soup = BeautifulSoup(html, 'html.parser')
            return soup.get_text().strip()
        except Exception as e:
            logger.warning(f"Failed to parse markdown, using raw content: {str(e)}")
            return content.strip()
    
    @staticmethod
    def extract_from_html(content: str) -> str:
        """Extract text from HTML"""
        try:
            soup = BeautifulSoup(content, 'html.parser')
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            return soup.get_text().strip()
        except Exception as e:
            logger.error(f"Failed to extract text from HTML: {str(e)}")
            return ""
    
    @staticmethod
    def extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to extract text from PDF {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX {file_path}: {str(e)}")
            return ""
    
    @staticmethod
    def extract_from_json(content: str) -> str:
        """Extract text from JSON"""
        try:
            data = json.loads(content)
            return TextExtractor._extract_text_from_dict(data)
        except Exception as e:
            logger.error(f"Failed to extract text from JSON: {str(e)}")
            return content
    
    @staticmethod
    def extract_from_yaml(content: str) -> str:
        """Extract text from YAML"""
        try:
            data = yaml.safe_load(content)
            return TextExtractor._extract_text_from_dict(data)
        except Exception as e:
            logger.error(f"Failed to extract text from YAML: {str(e)}")
            return content
    
    @staticmethod
    def _extract_text_from_dict(data: Any, prefix: str = "") -> str:
        """Recursively extract text from dictionary/list structures"""
        text_parts = []
        
        if isinstance(data, dict):
            for key, value in data.items():
                key_text = f"{prefix}{key}: " if prefix else f"{key}: "
                if isinstance(value, (dict, list)):
                    text_parts.append(key_text)
                    text_parts.append(TextExtractor._extract_text_from_dict(value, prefix + "  "))
                else:
                    text_parts.append(f"{key_text}{str(value)}")
        elif isinstance(data, list):
            for i, item in enumerate(data):
                if isinstance(item, (dict, list)):
                    text_parts.append(TextExtractor._extract_text_from_dict(item, prefix + "  "))
                else:
                    text_parts.append(f"{prefix}- {str(item)}")
        else:
            text_parts.append(str(data))
        
        return "\n".join(text_parts)


class TextChunker:
    """Chunks text into smaller pieces for processing"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def chunk_text(self, text: str, metadata: Optional[Dict[str, Any]] = None) -> List[DocumentChunk]:
        """
        Chunk text into smaller pieces.
        
        Args:
            text: Text to chunk
            metadata: Additional metadata for chunks
            
        Returns:
            List[DocumentChunk]: List of text chunks
        """
        if not text.strip():
            return []
        
        chunks = []
        metadata = metadata or {}
        
        # Try semantic chunking first (by paragraphs/sentences)
        semantic_chunks = self._semantic_chunk(text)
        
        if semantic_chunks:
            # Process semantic chunks
            for i, chunk_text in enumerate(semantic_chunks):
                if len(chunk_text.strip()) > 0:
                    chunk = DocumentChunk(
                        content=chunk_text.strip(),
                        chunk_index=i,
                        start_char=text.find(chunk_text),
                        end_char=text.find(chunk_text) + len(chunk_text),
                        metadata={**metadata, "chunk_type": "semantic"}
                    )
                    chunks.append(chunk)
        else:
            # Fall back to fixed-size chunking
            chunks = self._fixed_size_chunk(text, metadata)
        
        return chunks
    
    def _semantic_chunk(self, text: str) -> List[str]:
        """Chunk text by semantic boundaries (paragraphs, sentences)"""
        try:
            # Split by double newlines (paragraphs)
            paragraphs = re.split(r'\n\s*\n', text)
            
            chunks = []
            current_chunk = ""
            
            for paragraph in paragraphs:
                paragraph = paragraph.strip()
                if not paragraph:
                    continue
                
                # If adding this paragraph would exceed chunk size
                if len(current_chunk) + len(paragraph) > self.chunk_size:
                    if current_chunk:
                        chunks.append(current_chunk)
                        # Start new chunk with overlap
                        overlap_text = self._get_overlap_text(current_chunk)
                        current_chunk = overlap_text + paragraph
                    else:
                        # Paragraph is too long, split by sentences
                        sentence_chunks = self._split_by_sentences(paragraph)
                        chunks.extend(sentence_chunks)
                else:
                    current_chunk += ("\n\n" if current_chunk else "") + paragraph
            
            # Add final chunk
            if current_chunk:
                chunks.append(current_chunk)
            
            return chunks
            
        except Exception as e:
            logger.warning(f"Semantic chunking failed: {str(e)}")
            return []
    
    def _split_by_sentences(self, text: str) -> List[str]:
        """Split long text by sentences"""
        # Simple sentence splitting
        sentences = re.split(r'[.!?]+\s+', text)
        
        chunks = []
        current_chunk = ""
        
        for sentence in sentences:
            sentence = sentence.strip()
            if not sentence:
                continue
            
            if len(current_chunk) + len(sentence) > self.chunk_size:
                if current_chunk:
                    chunks.append(current_chunk)
                    current_chunk = sentence
                else:
                    # Sentence is too long, use fixed-size chunking
                    chunks.extend(self._fixed_size_chunk(sentence, {}))
            else:
                current_chunk += (". " if current_chunk else "") + sentence
        
        if current_chunk:
            chunks.append(current_chunk)
        
        return chunks
    
    def _fixed_size_chunk(self, text: str, metadata: Dict[str, Any]) -> List[DocumentChunk]:
        """Chunk text using fixed-size windows"""
        chunks = []
        start = 0
        chunk_index = 0
        
        while start < len(text):
            end = min(start + self.chunk_size, len(text))
            
            # Try to end at word boundary
            if end < len(text):
                last_space = text.rfind(' ', start, end)
                if last_space > start:
                    end = last_space
            
            chunk_text = text[start:end].strip()
            
            if chunk_text:
                chunk = DocumentChunk(
                    content=chunk_text,
                    chunk_index=chunk_index,
                    start_char=start,
                    end_char=end,
                    metadata={**metadata, "chunk_type": "fixed_size"}
                )
                chunks.append(chunk)
                chunk_index += 1
            
            # Move start position with overlap
            start = max(start + self.chunk_size - self.chunk_overlap, end)
        
        return chunks
    
    def _get_overlap_text(self, text: str) -> str:
        """Get overlap text from the end of current chunk"""
        if len(text) <= self.chunk_overlap:
            return text
        
        overlap_start = len(text) - self.chunk_overlap
        # Try to start at word boundary
        space_pos = text.find(' ', overlap_start)
        if space_pos != -1:
            overlap_start = space_pos + 1
        
        return text[overlap_start:] + "\n\n"


class DocumentProcessor:
    """
    Main document processor for RAG system.
    
    Handles document ingestion, text extraction, chunking,
    and preparation for vector storage.
    """
    
    def __init__(self, settings: Settings):
        self.settings = settings
        self.text_extractor = TextExtractor()
        self.text_chunker = TextChunker(
            chunk_size=settings.vector.CHUNK_SIZE,
            chunk_overlap=settings.vector.CHUNK_OVERLAP
        )
        self.embedding_service = get_embedding_service(settings)
    
    def detect_document_type(self, file_path: Optional[str] = None, 
                           content: Optional[str] = None) -> DocumentType:
        """Detect document type from file path or content"""
        if file_path:
            path = Path(file_path)
            extension = path.suffix.lower()
            
            type_mapping = {
                '.txt': DocumentType.TEXT,
                '.md': DocumentType.MARKDOWN,
                '.markdown': DocumentType.MARKDOWN,
                '.pdf': DocumentType.PDF,
                '.docx': DocumentType.DOCX,
                '.doc': DocumentType.DOCX,
                '.html': DocumentType.HTML,
                '.htm': DocumentType.HTML,
                '.json': DocumentType.JSON,
                '.yaml': DocumentType.YAML,
                '.yml': DocumentType.YAML,
                '.py': DocumentType.CODE,
                '.js': DocumentType.CODE,
                '.ts': DocumentType.CODE,
                '.java': DocumentType.CODE,
                '.cpp': DocumentType.CODE,
                '.c': DocumentType.CODE,
            }
            
            return type_mapping.get(extension, DocumentType.UNKNOWN)
        
        elif content:
            # Try to detect from content
            content_lower = content.lower().strip()
            
            if content_lower.startswith('<!doctype html') or content_lower.startswith('<html'):
                return DocumentType.HTML
            elif content_lower.startswith('{') and content_lower.endswith('}'):
                try:
                    json.loads(content)
                    return DocumentType.JSON
                except:
                    pass
            elif content_lower.startswith('---') or re.match(r'^[a-zA-Z_][a-zA-Z0-9_]*:', content):
                try:
                    yaml.safe_load(content)
                    return DocumentType.YAML
                except:
                    pass
            elif re.search(r'#+\s+', content) or re.search(r'\*\*.*\*\*', content):
                return DocumentType.MARKDOWN
        
        return DocumentType.TEXT
    
    async def process_file(self, file_path: str, 
                          title: Optional[str] = None,
                          metadata: Optional[Dict[str, Any]] = None) -> ProcessedDocument:
        """
        Process a document file.
        
        Args:
            file_path: Path to the document file
            title: Optional document title
            metadata: Additional metadata
            
        Returns:
            ProcessedDocument: Processed document with chunks
        """
        try:
            path = Path(file_path)
            if not path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Detect document type
            doc_type = self.detect_document_type(file_path=file_path)
            
            # Extract text based on document type
            if doc_type == DocumentType.PDF:
                content = self.text_extractor.extract_from_pdf(file_path)
            elif doc_type == DocumentType.DOCX:
                content = self.text_extractor.extract_from_docx(file_path)
            else:
                # Read as text file
                with open(file_path, 'r', encoding='utf-8') as f:
                    raw_content = f.read()
                
                if doc_type == DocumentType.MARKDOWN:
                    content = self.text_extractor.extract_from_markdown(raw_content)
                elif doc_type == DocumentType.HTML:
                    content = self.text_extractor.extract_from_html(raw_content)
                elif doc_type == DocumentType.JSON:
                    content = self.text_extractor.extract_from_json(raw_content)
                elif doc_type == DocumentType.YAML:
                    content = self.text_extractor.extract_from_yaml(raw_content)
                else:
                    content = self.text_extractor.extract_from_text(raw_content)
            
            # Create processed document
            document = ProcessedDocument(
                title=title or path.stem,
                content=content,
                document_type=doc_type,
                source_path=file_path,
                metadata={
                    "file_size": path.stat().st_size,
                    "file_extension": path.suffix,
                    "file_name": path.name,
                    **(metadata or {})
                }
            )
            
            # Process chunks and embeddings
            await self._process_document_chunks(document)
            
            logger.info(f"Processed document: {file_path} ({len(document.chunks)} chunks)")
            return document
            
        except Exception as e:
            logger.error(f"Failed to process file {file_path}: {str(e)}")
            raise
    
    async def process_text(self, content: str,
                          title: str,
                          source_url: Optional[str] = None,
                          metadata: Optional[Dict[str, Any]] = None) -> ProcessedDocument:
        """
        Process text content directly.
        
        Args:
            content: Text content to process
            title: Document title
            source_url: Optional source URL
            metadata: Additional metadata
            
        Returns:
            ProcessedDocument: Processed document with chunks
        """
        try:
            # Detect document type from content
            doc_type = self.detect_document_type(content=content)
            
            # Extract text based on detected type
            if doc_type == DocumentType.MARKDOWN:
                processed_content = self.text_extractor.extract_from_markdown(content)
            elif doc_type == DocumentType.HTML:
                processed_content = self.text_extractor.extract_from_html(content)
            elif doc_type == DocumentType.JSON:
                processed_content = self.text_extractor.extract_from_json(content)
            elif doc_type == DocumentType.YAML:
                processed_content = self.text_extractor.extract_from_yaml(content)
            else:
                processed_content = self.text_extractor.extract_from_text(content)
            
            # Create processed document
            document = ProcessedDocument(
                title=title,
                content=processed_content,
                document_type=doc_type,
                source_url=source_url,
                metadata={
                    "content_length": len(content),
                    "processed_length": len(processed_content),
                    **(metadata or {})
                }
            )
            
            # Process chunks and embeddings
            await self._process_document_chunks(document)
            
            logger.info(f"Processed text document: {title} ({len(document.chunks)} chunks)")
            return document
            
        except Exception as e:
            logger.error(f"Failed to process text content: {str(e)}")
            raise
    
    async def _process_document_chunks(self, document: ProcessedDocument) -> None:
        """Process document chunks and generate embeddings"""
        try:
            # Generate chunks
            chunks = self.text_chunker.chunk_text(
                document.content,
                metadata={
                    "document_id": document.id,
                    "document_title": document.title,
                    "document_type": document.document_type.value
                }
            )
            
            # Generate embeddings for chunks
            if chunks:
                chunk_texts = [chunk.content for chunk in chunks]
                embedding_results = await self.embedding_service.generate_embeddings(chunk_texts)
                
                for chunk, embedding_result in zip(chunks, embedding_results):
                    chunk.embedding = embedding_result.embedding
            
            # Generate document-level embedding (from first chunk or summary)
            if chunks:
                document.embedding = chunks[0].embedding
            
            document.chunks = chunks
            
        except Exception as e:
            logger.error(f"Failed to process document chunks: {str(e)}")
            raise
    
    async def batch_process_files(self, file_paths: List[str],
                                 metadata: Optional[Dict[str, Any]] = None) -> List[ProcessedDocument]:
        """
        Process multiple files in batch.
        
        Args:
            file_paths: List of file paths to process
            metadata: Common metadata for all files
            
        Returns:
            List[ProcessedDocument]: List of processed documents
        """
        documents = []
        
        # Process files concurrently
        tasks = []
        for file_path in file_paths:
            task = self.process_file(file_path, metadata=metadata)
            tasks.append(task)
        
        # Wait for all tasks to complete
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                logger.error(f"Failed to process {file_paths[i]}: {str(result)}")
            else:
                documents.append(result)
        
        logger.info(f"Batch processed {len(documents)}/{len(file_paths)} files successfully")
        return documents
    
    def get_supported_extensions(self) -> List[str]:
        """Get list of supported file extensions"""
        return ['.txt', '.md', '.markdown', '.pdf', '.docx', '.doc', 
                '.html', '.htm', '.json', '.yaml', '.yml', '.py', 
                '.js', '.ts', '.java', '.cpp', '.c']
